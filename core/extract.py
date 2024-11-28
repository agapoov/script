from docx import Document


def extract_nested_table_from_parent(file_path):
    doc = Document(file_path)
    nested_tables_data = []
    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[2]
            if cell.tables:
                for nested_table in cell.tables:
                    nested_table_data = [
                        [cell.text.strip() for cell in row.cells]
                        for idx, row in enumerate(nested_table.rows)
                        if idx > 0
                    ]
                    nested_tables_data.append(nested_table_data)

    return nested_tables_data
