from docx import Document


def save_updated_data_to_word(updated_tables, file_path):
    doc = Document()

    for table_index, table in enumerate(updated_tables):
        doc.add_paragraph(f"Таблица {table_index + 1}:")

        table_in_doc = doc.add_table(rows=1, cols=6)
        table_in_doc.style = 'Table Grid'
        hdr_cells = table_in_doc.rows[0].cells
        hdr_cells[0].text = '№ п/п'
        hdr_cells[1].text = 'Групповой показатель'
        hdr_cells[2].text = 'Средство проверки'
        hdr_cells[3].text = 'Эталонное значение'
        hdr_cells[4].text = 'Да / 1'
        hdr_cells[5].text = 'Нет / 0'

        prev_command = None
        prev_row = None

        for row_idx, row in enumerate(table):
            row_cells = table_in_doc.add_row().cells

            current_command = row[2]

            row_cells[0].text = str(row[0])
            row_cells[1].text = row[1]
            row_cells[2].text = current_command
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]

            if current_command == prev_command and prev_row is not None:
                row_cells[2].text = ""
                table_in_doc.cell(row_idx, 2).merge(prev_row)

            else:
                prev_row = row_cells[2]

            prev_command = current_command

        if prev_row is not None and prev_row != row_cells[2]:
            prev_row.merge(row_cells[2])

        doc.add_paragraph()

    doc.save(file_path)
